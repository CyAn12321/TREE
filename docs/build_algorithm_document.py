from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "TREE_项目参考与算法设计说明.md"
OUTPUT = ROOT / "TREE_项目参考与预期算法设计说明.docx"
BUILD_DIR = ROOT / ".build"
DIAGRAM = BUILD_DIR / "architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "C8D2DD"

LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"


def set_run_font(run, size=None, color=None, bold=None, italic=None, mono=False):
    font_name = MONO_FONT if mono else LATIN_FONT
    east_asia = MONO_FONT if mono else CJK_FONT
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn("w:" + margin_name))
        if node is None:
            node = OxmlElement("w:" + margin_name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn("w:" + edge))
        if tag is None:
            tag = OxmlElement("w:" + edge)
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440.0)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.extend((r_fonts, color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://[^\s；;）)]+)")


def add_rich_text(paragraph, text, size=11, color=None):
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(9, size - 0.5), color=NAVY, mono=True)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "EEF2F6")
            run._element.get_or_add_rPr().append(shade)
        else:
            add_hyperlink(paragraph, token, token)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def paragraph_bottom_border(paragraph, color="9CB4C8", size=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [style.name for style in styles]:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = MONO_FONT
    code._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("TREE  |  参数化植物建模与纯顶点动画")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    paragraph_bottom_border(p, color="CBD5DF", size=6)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run("算法设计基线  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)


def draw_architecture(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 920
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_bold = ImageFont.truetype(r"C:\Windows\Fonts\msyhbd.ttc", 42)
    font = ImageFont.truetype(r"C:\Windows\Fonts\msyh.ttc", 31)
    font_small = ImageFont.truetype(r"C:\Windows\Fonts\msyh.ttc", 25)

    boxes = [
        ("参数与预设", "树种 / 层级 / 角度 / 种子"),
        ("L-System 推导", "Typed L-string + Branch Graph"),
        ("3D Turtle 解释", "枝干网格 + 附着点 + 顶点属性"),
        ("器官与季节", "资产筛选 / 状态机 / 逻辑实例"),
        ("纯顶点天气动画", "风 / 脱落 / 雨雪 / 积雪"),
    ]
    x0, x1 = 190, 1410
    y = 54
    box_h = 126
    gap = 48
    colors = ["#DCE9F2", "#CFE2F3", "#D9EAD3", "#FFF2CC", "#FCE5CD"]
    for index, ((title, subtitle), fill) in enumerate(zip(boxes, colors)):
        draw.rounded_rectangle(
            (x0, y, x1, y + box_h),
            radius=28,
            fill=fill,
            outline="#6B879C",
            width=4,
        )
        draw.text((x0 + 42, y + 22), title, font=font_bold, fill="#17324D")
        draw.text((x0 + 510, y + 37), subtitle, font=font, fill="#344A5E")
        if index < len(boxes) - 1:
            center_x = (x0 + x1) // 2
            arrow_top = y + box_h + 7
            arrow_bottom = y + box_h + gap - 8
            draw.line((center_x, arrow_top, center_x, arrow_bottom), fill="#52758F", width=7)
            draw.polygon(
                [
                    (center_x - 16, arrow_bottom - 14),
                    (center_x + 16, arrow_bottom - 14),
                    (center_x, arrow_bottom + 10),
                ],
                fill="#52758F",
            )
        y += box_h + gap

    draw.text(
        (190, 885),
        "固定拓扑：生成阶段创建全部顶点；动画阶段只更新顶点位置。",
        font=font_small,
        fill="#667085",
    )
    image.save(path, dpi=(180, 180))


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("TREE PROJECT")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("参考依据与预期算法设计说明")
    set_run_font(run, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("L-System 树干 · 器官实例与季节系统 · 纯顶点天气动画")
    set_run_font(run, size=14, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(26)
    paragraph_bottom_border(rule, color="7FA0B7", size=16)

    rows = [
        ("文档版本", "v1.0（算法设计基线）"),
        ("日期", "2026-07-20"),
        ("目标环境", "Autodesk Maya / Python；可迁移到实时 vertex shader"),
        ("核心约束", "天气动画固定拓扑，仅更新顶点位置；不使用骨骼、刚体或粒子求解器"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table, color="D8E0E7", size=4)
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], LIGHT_BLUE)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=10.5, color=NAVY, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_rich_text(p1, value, size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用于小组中期汇报、后续实现分工与验收")
    set_run_font(run, size=10, color=MUTED, italic=True)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_navigation(doc):
    p = doc.add_paragraph("内容导览", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    entries = [
        "0. 文档结论",
        "1. 参考来源与本项目采用内容",
        "2. 总体数据流",
        "3. 树干生成算法",
        "4. 附着点输出与器官资产系统",
        "5. 季节系统",
        "6. 纯顶点天气动画",
        "7. 模块接口建议",
        "8. 实现顺序",
        "9. 验收标准",
        "10. 采用边界、风险与说明",
        "参考文献与项目",
    ]
    for entry in entries:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, entry)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(8)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT_GRAY)
    note._p.get_or_add_pPr().append(shade)
    add_rich_text(
        note,
        "阅读提示：正文中的 [n] 与末尾参考文献编号对应；“采用”表示转化为本项目设计，“不采用”表示未复现完整生物或物理模型。",
        size=10.5,
        color=NAVY,
    )
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def widths_for_table(column_count):
    if column_count == 3:
        return [1750, 2900, 4710]
    if column_count == 6:
        return [800, 1100, 1900, 1100, 2100, 2360]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths_for_table(column_count))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(rows):
        for col_index in range(column_count):
            cell = table.rows[row_index].cells[col_index]
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            value = row_values[col_index] if col_index < len(row_values) else ""
            add_rich_text(p, value, size=9.2 if column_count >= 6 else 9.7)
            for run in p.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_body(doc, source_text):
    lines = source_text.splitlines()
    # Skip the Markdown title/metadata already represented by the cover.
    start = next(index for index, line in enumerate(lines) if line.startswith("## 0."))
    lines = lines[start:]
    index = 0
    in_code = False
    code_lines = []
    paragraph_buffer = []
    inserted_diagram = False

    def flush_paragraph():
        if not paragraph_buffer:
            return
        source = " ".join(part.strip() for part in paragraph_buffer)
        p = doc.add_paragraph()
        add_rich_text(p, source)
        # A standalone formula should travel with the explanation that follows it;
        # otherwise Word may leave the formula sitting directly on the footer.
        if source.startswith("`") and source.endswith("`"):
            p.paragraph_format.keep_with_next = True
        paragraph_buffer[:] = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_GRAY)
                p_pr.append(shd)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=9, color=NAVY, mono=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.lstrip().startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1)) - 1
            text = heading.group(2)
            p = doc.add_paragraph(style="Heading {}".format(min(level, 3)))
            add_rich_text(p, text, size={1: 16, 2: 13, 3: 12}[min(level, 3)], color=BLUE if level < 3 else DARK_BLUE)
            if text == "2. 总体数据流" and not inserted_diagram:
                picture = doc.add_paragraph()
                picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture.paragraph_format.space_after = Pt(4)
                inline_shape = picture.add_run().add_picture(
                    str(DIAGRAM), width=Inches(6.15)
                )
                inline_shape._inline.docPr.set(
                    "descr", "TREE project algorithm architecture flowchart"
                )
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                run = caption.add_run("图 1  TREE 项目算法数据流与模块边界")
                set_run_font(run, size=9, color=MUTED, italic=True)
                inserted_diagram = True
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph()
            if bullet:
                p = doc.add_paragraph(style="List Bullet")
                add_rich_text(p, bullet.group(1))
            else:
                # Keep the explicit Markdown number. Word's built-in List Number
                # style otherwise continues numbering across unrelated sections.
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.25
                add_rich_text(p, line)
            index += 1
            continue

        if not line.strip():
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def audit_document(doc):
    assert len(doc.sections) == 1
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert all(
        round(value.inches, 2) == 1.0
        for value in (
            section.top_margin,
            section.bottom_margin,
            section.left_margin,
            section.right_margin,
        )
    )
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        assert doc.styles[style_name].font.name == LATIN_FONT
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa"
        assert int(tbl_w.get(qn("w:w"))) == 9360


def main():
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    draw_architecture(DIAGRAM)
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_cover(doc)
    add_navigation(doc)
    add_body(doc, source_text)
    audit_document(doc)
    doc.core_properties.title = "TREE 项目：参考依据与预期算法设计说明"
    doc.core_properties.subject = "L-System、季节器官实例与纯顶点天气动画"
    doc.core_properties.author = "TREE 项目组"
    doc.core_properties.keywords = "L-System, LPy, PlantGL, Maya, season, vertex animation"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
